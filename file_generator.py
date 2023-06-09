import os

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from random import randint, sample
from pprint import pprint
import pymorphy2
from docx2pdf import convert


def system_translation(number, radixresult):
    digits = "0123456789ABCDEF"
    res = ''
    number = abs(number)
    while number != 0:
        res = str(digits[number % radixresult]) + res
        number //= radixresult
    return res


def generate_exercise(numbertasks, operation, radix):
    res = [(system_translation(randint(20, 512), radix),
           system_translation(randint(20, 512), radix)) for i in range(numbertasks)]
    resanswers = [system_translation(eval(str(int(item[0], radix)) + operation +
                       str(int(item[1], radix))), radix) for item in res]
    return res, resanswers

def generate_ariphmetic_txt(theme, number_of_exercises, operations, radixs, fileformat, filedirectory, variant_number):
    f = open(filedirectory + "/" + theme + " Вариант " + str(variant_number) + '.txt', 'w')

    f.write(theme + "\n")
    f.write("Вариант " + str(variant_number) + "\n")
    points = 'абвгдежзиклмн'
    for i in range(number_of_exercises):
        f.write('№ ' + str(i + 1) + ". Выполните арифметические операции в " + str(radixs[i % 4]) + " с.с.:" + "\n")
        exercises = generate_exercise(6, operations[i % 3], radixs[i % 4])
        for item, point in zip(exercises[0], points):
            f.write("\t" + point + ") " + item[0] + " " + operations[i % 3] + " "+ item[1])
            f.write("\n")
    f.close()


def generate_systems_txt(number_of_exercises, radixfrom, fileformat, filedirectory, variant_number):
    f = open(filedirectory + "/" + "СР Перевод между системами счисления" + " Вариант "
             + str(variant_number) + '.txt', 'w')

    f.write("Cамостоятельная работа" + "\n")
    f.write("Перевод между системами счисления" + "\n")

    points = 'абвгдежзиклмн'

    morph = pymorphy2.MorphAnalyzer()

    d = {
        "двоичная": 2,
        "четверичная": 4,
        "восьмеричная": 8,
        "шестнадцатеричная": 16
    }

    for exer in range(number_of_exercises):
        systems = sample(radixfrom, 2)

        word1 = morph.parse(systems[0])[0].inflect({"ADJF", "femn", "loct"})[0]
        word2 = morph.parse(systems[1])[0].inflect({"ADJF", "femn", "accs"})[0]
        f.write('№ ' + str(exer + 1) + ". Выполните перевод из " +
                                word1 + " в " + word2 + "\n")
        exercises = [system_translation(randint(10, 1000), d[systems[0]]) for i in range(4)]
        pprint(exercises)
        for item, point in zip(exercises, points):
            f.write("\t" + point + ") " + item + "\n")

    f.close()



def generate_ariphmetic_pdf(theme, number_of_exercises, operations, radixs, fileformat, filedirectory, variant_number):
    print("сюда дошло")
    answers_document = Document()
    answers_nextp = answers_document.add_paragraph()
    answers_nextp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    answers_nextrun = answers_nextp.add_run("Ответы")
    answers_nextrun.font.size = Pt(16)
    answers_nextrun.font.bold = True
    answers_nextrun.font.name = 'Times New Roman'

    document = Document()
    nextp = document.add_paragraph()
    nextp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nextrun = nextp.add_run("Cамостоятельная работа")
    nextrun.font.size = Pt(16)
    nextrun.font.bold = True
    nextrun.font.name = 'Times New Roman'

    nextp = document.add_paragraph()
    nextp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nextrun = nextp.add_run(theme)
    nextrun.font.size = Pt(16)
    nextrun.font.bold = True
    nextrun.font.name = 'Times New Roman'
    points = 'абвгдежзиклмн'
    for i in range(number_of_exercises):
        nextp = document.add_paragraph()
        nextp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        nextrun = nextp.add_run('№ ' + str(i + 1) + ". Выполните арифметические операции:")
        nextrun.font.size = Pt(14)
        nextrun.font.bold = False
        nextrun.font.name = 'Times New Roman'

        answers_nextp = answers_document.add_paragraph()
        answers_nextp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        answers_nextrun = answers_nextp.add_run('№ ' + str(i + 1) + ".")
        answers_nextrun.font.size = Pt(14)
        answers_nextrun.font.bold = False
        answers_nextrun.font.name = 'Times New Roman'
        exercises = generate_exercise(6, operations[i % 3], radixs[i % 4])
        for item, point in zip(exercises[0], points):
            nextp = document.add_paragraph()
            nextp.paragraph_format.first_line_indent = Inches(0.5)
            nextp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            nextrun = nextp.add_run(point + ") " + item[0])
            nextrun.subscript = False
            nextrun.font.size = Pt(14)
            nextrun.font.bold = False
            nextrun.font.name = 'Times New Roman'
            subnextrun = nextp.add_run(str(radixs[i % 4 ]))
            subnextrun.subscript = True
            nextrun = nextp.add_run(f" {operations[i % 3]} " + item[1])
            subnextrun = nextp.add_run(str(radixs[i % 4]))
            subnextrun.subscript = True
            nextrun.font.size = Pt(14)
            nextrun.font.bold = False
            nextrun.font.name = 'Times New Roman'

        for item, point in zip(exercises[1], points):
            answers_nextp = answers_document.add_paragraph()
            answers_nextp.paragraph_format.first_line_indent = Inches(0.5)
            answers_nextp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            answers_nextrun = answers_nextp.add_run(point + ") " + item)
            answers_nextrun.subscript = False
            answers_nextrun.font.size = Pt(14)
            answers_nextrun.font.bold = False
            answers_nextrun.font.name = 'Times New Roman'
            answers_subnextrun = answers_nextp.add_run(str(radixs[i % 4]))
            answers_subnextrun.subscript = True

    if fileformat == ".docx":
        document.save(filedirectory + "/" + theme + " Вариант " + str(variant_number + 1) + '.docx')
        answers_document.save(filedirectory + "/" + theme + " Вариант " + str(variant_number + 1) + " ОТВЕТЫ.docx")
    else:
        document.save(filedirectory + "/" + theme + " Вариант " + str(variant_number + 1) + '.docx')
        answers_document.save(filedirectory + "/" + theme + " Вариант " + str(variant_number + 1) + " ОТВЕТЫ.docx")
        convert(filedirectory + "/" + theme + " Вариант " + str(variant_number + 1) + '.docx')
        convert(filedirectory + "/" + theme + " Вариант " + str(variant_number + 1) + " ОТВЕТЫ.docx")
        os.remove(filedirectory + "/" + theme + " Вариант " + str(variant_number + 1) + '.docx')
        os.remove(filedirectory + "/" + theme + " Вариант " + str(variant_number + 1) + " ОТВЕТЫ.docx")


def generate_systems_pdf(number_of_exercises, radixfrom, fileformat, filedirectory, variant_number):
    document = Document()
    nextp = document.add_paragraph()
    nextp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nextrun = nextp.add_run("Cамостоятельная работа")
    nextrun.font.size = Pt(16)
    nextrun.font.bold = True
    nextrun.font.name = 'Times New Roman'

    nextp = document.add_paragraph()
    nextp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nextrun = nextp.add_run("Перевод между системами счисления")
    nextrun.font.size = Pt(16)
    nextrun.font.bold = True
    nextrun.font.name = 'Times New Roman'
    points = 'абвгдежзиклмн'

    morph = pymorphy2.MorphAnalyzer()

    d = {
        "двоичная": 2,
        "четверичная": 4,
        "восьмеричная": 8,
        "шестнадцатеричная": 16
    }

    for exer in range(number_of_exercises):
        nextp = document.add_paragraph()
        nextp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        systems = sample(radixfrom, 2)

        word1 = morph.parse(systems[0])[0].inflect({"ADJF", "femn", "loct"})[0]
        word2 = morph.parse(systems[1])[0].inflect({"ADJF", "femn", "accs"})[0]
        nextrun = nextp.add_run('№ ' + str(exer + 1) + ". Выполните перевод из " +
                                word1 + " в " + word2)
        nextrun.font.size = Pt(14)
        nextrun.font.bold = False
        nextrun.font.name = 'Times New Roman'
        exercises = [system_translation(randint(10, 1000), d[systems[0]]) for i in range(4)]
        pprint(exercises)
        for item, point in zip(exercises, points):
            nextp = document.add_paragraph()
            nextp.paragraph_format.first_line_indent = Inches(0.5)
            nextp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            nextrun = nextp.add_run(point + ") " + item)
            nextrun.subscript = False
            nextrun.font.size = Pt(14)
            nextrun.font.bold = False
            nextrun.font.name = 'Times New Roman'
            subnextrun = nextp.add_run(str(d[systems[0]]))
            subnextrun.subscript = True
            nextrun.font.size = Pt(14)
            nextrun.font.bold = False
            nextrun.font.name = 'Times New Roman'

    if fileformat == ".docx":
        document.save(filedirectory + "/" + 'СР Перевод между системами счисления' +
                      " Вариант " + str(variant_number) + '.docx')
    else:
        document.save(
            filedirectory + "/" + 'СР Перевод между системами счисления' +
            " Вариант " + str(variant_number + 1) + '.docx')
        convert(filedirectory + "/" + 'СР Перевод между системами счисления' +
                " Вариант " + str(variant_number + 1) + '.docx')
        os.remove(filedirectory + "/" + 'СР Перевод между системами счисления' +
                  " Вариант " + str(variant_number + 1) + '.docx')