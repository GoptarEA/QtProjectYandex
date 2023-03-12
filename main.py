from PyQt6.QtWidgets import QApplication, QMainWindow, QStackedWidget, QDialog
import sys
from ui.startWindow import Ui_StartWindow
from ui.loginWindow import Ui_LoginWindow
from ui.registerWindow import Ui_RegisterWindow
from ui.roleWindow import Ui_RoleWindow
from ui.main_menu import Ui_MainMenu
from ui.dialog_win import Ui_Dialog
from ui.generate_test import Ui_GenerateTestWindow
import sqlite3
import hashlib
import os
import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from random import randint, sample
from pprint import pprint
import pymorphy2
from docx2pdf import convert
import base64


def system_translation(number, radixresult):
    digits = "0123456789ABCDEF"
    res = ''
    number = abs(number)
    while number != 0:
        res = str(digits[number % radixresult]) + res
        number //= radixresult
    return res


def generate_exercise(numbertasks, operation, radix):
    res = [(system_translation(randint(20, 512), radix),
           system_translation(randint(20, 512), radix)) for i in range(numbertasks)]
    resanswers = [system_translation(eval(str(int(item[0], radix)) + operation +
                       str(int(item[1], radix))), radix) for item in res]
    return res, resanswers


def generate_ariphmetic_pdf(theme, number_of_exercises, operations, radixs):

    answers_document = Document()
    answers_nextp = answers_document.add_paragraph()
    answers_nextp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    answers_nextrun = answers_nextp.add_run("Ответы")
    answers_nextrun.font.size = Pt(16)
    answers_nextrun.font.bold = True
    answers_nextrun.font.name = 'Times New Roman'

    document = Document()
    nextp = document.add_paragraph()
    nextp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nextrun = nextp.add_run("Cамостоятельная работа")
    nextrun.font.size = Pt(16)
    nextrun.font.bold = True
    nextrun.font.name = 'Times New Roman'

    nextp = document.add_paragraph()
    nextp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nextrun = nextp.add_run(theme)
    nextrun.font.size = Pt(16)
    nextrun.font.bold = True
    nextrun.font.name = 'Times New Roman'
    points = 'абвгдежзиклмн'
    for i in range(number_of_exercises):
        nextp = document.add_paragraph()
        nextp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        nextrun = nextp.add_run('№ ' + str(i + 1) + ". Выполните арифметические операции:")
        nextrun.font.size = Pt(14)
        nextrun.font.bold = False
        nextrun.font.name = 'Times New Roman'

        answers_nextp = answers_document.add_paragraph()
        answers_nextp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        answers_nextrun = answers_nextp.add_run('№ ' + str(i + 1) + ".")
        answers_nextrun.font.size = Pt(14)
        answers_nextrun.font.bold = False
        answers_nextrun.font.name = 'Times New Roman'
        exercises = generate_exercise(6, operations[i % 3], radixs[i % 4])
        for item, point in zip(exercises[0], points):
            nextp = document.add_paragraph()
            nextp.paragraph_format.first_line_indent = Inches(0.5)
            nextp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            nextrun = nextp.add_run(point + ") " + item[0])
            nextrun.subscript = False
            nextrun.font.size = Pt(14)
            nextrun.font.bold = False
            nextrun.font.name = 'Times New Roman'
            subnextrun = nextp.add_run(str(radixs[i % 4 ]))
            subnextrun.subscript = True
            nextrun = nextp.add_run(f" {operations[i % 3]} " + item[1])
            subnextrun = nextp.add_run(str(radixs[i % 4]))
            subnextrun.subscript = True
            nextrun.font.size = Pt(14)
            nextrun.font.bold = False
            nextrun.font.name = 'Times New Roman'

        for item, point in zip(exercises[1], points):
            answers_nextp = answers_document.add_paragraph()
            answers_nextp.paragraph_format.first_line_indent = Inches(0.5)
            answers_nextp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            answers_nextrun = answers_nextp.add_run(point + ") " + item)
            answers_nextrun.subscript = False
            answers_nextrun.font.size = Pt(14)
            answers_nextrun.font.bold = False
            answers_nextrun.font.name = 'Times New Roman'
            answers_subnextrun = answers_nextp.add_run(str(radixs[i % 4 ]))
            answers_subnextrun.subscript = True

    document.save(theme + '.docx')
    convert(theme + '.docx')

    answers_document.save(theme + " ОТВЕТЫ.docx")
    convert(theme + " ОТВЕТЫ.docx")


def generate_systems_pdf(number_of_exercises, radixfrom):
    document = Document()
    nextp = document.add_paragraph()
    nextp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nextrun = nextp.add_run("Cамостоятельная работа")
    nextrun.font.size = Pt(16)
    nextrun.font.bold = True
    nextrun.font.name = 'Times New Roman'

    nextp = document.add_paragraph()
    nextp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nextrun = nextp.add_run("Перевод между системами счисления")
    nextrun.font.size = Pt(16)
    nextrun.font.bold = True
    nextrun.font.name = 'Times New Roman'
    points = 'абвгдежзиклмн'

    morph = pymorphy2.MorphAnalyzer()

    d = {
        "двоичная": 2,
        "четверичная": 4,
        "восьмеричная": 8,
        "шестнадцатеричная": 16
    }

    for exer in range(number_of_exercises):
        nextp = document.add_paragraph()
        nextp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        systems = sample(radixfrom, 2)

        word1 = morph.parse(systems[0])[0].inflect({"ADJF", "femn", "loct"})[0]
        word2 = morph.parse(systems[1])[0].inflect({"ADJF", "femn", "accs"})[0]
        nextrun = nextp.add_run('№ ' + str(exer + 1) + ". Выполните перевод из " +
                                word1 + " в " + word2)
        nextrun.font.size = Pt(14)
        nextrun.font.bold = False
        nextrun.font.name = 'Times New Roman'
        exercises = [system_translation(randint(10, 1000), d[systems[0]]) for i in range(4)]
        pprint(exercises)
        for item, point in zip(exercises, points):
            nextp = document.add_paragraph()
            nextp.paragraph_format.first_line_indent = Inches(0.5)
            nextp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            nextrun = nextp.add_run(point + ") " + item)
            nextrun.subscript = False
            nextrun.font.size = Pt(14)
            nextrun.font.bold = False
            nextrun.font.name = 'Times New Roman'
            subnextrun = nextp.add_run(str(d[systems[0]]))
            subnextrun.subscript = True
            nextrun.font.size = Pt(14)
            nextrun.font.bold = False
            nextrun.font.name = 'Times New Roman'

    document.save('СР Перевод между системами счисления.docx')
    convert('СР Перевод между системами счисления.docx')


users = sqlite3.connect('data/users.db')
users_cur = users.cursor()

works = sqlite3.connect('data/works.db')
works_cur = works.cursor()


def create_db():
    users_cur.execute("""CREATE TABLE IF NOT EXISTS users(
       userid INTEGER PRIMARY KEY AUTOINCREMENT,
       login TEXT NOT NULL,
       password TEXT NOT NULL);
    """)
    users.commit()

    works_cur.execute("""CREATE TABLE IF NOT EXISTS works(
           workid INTEGER PRIMARY KEY AUTOINCREMENT,
           theme TEXT NOT NULL,
           varcount INTEGER NOT NULL,
           excount INTEGER NOT NULL,
           creation_date TEXT NOT NULL);
        """)
    works.commit()


def add_new_user(new_login, new_password):
    s = base64.b64encode(new_password).decode('utf-8')
    users_cur.execute(f"INSERT INTO users (login, password) VALUES ('{new_login}', '{s}');")
    users.commit()


def add_new_work(theme, varcount, excount):
    works_cur.execute("INSERT INTO works VALUES(?, ?, ?, ?);", (theme, varcount, excount, datetime.date.today()))


def check_user(login):
    print(login)
    users_cur.execute(f"SELECT * FROM users;")
    users_list = users_cur.fetchall()
    print(users_list)
    return any([login in elem for elem in users_list])


def get_user(login):
    print(login)
    users_cur.execute(f"SELECT * FROM users;")
    users_list = users_cur.fetchall()
    print('ok')
    return [elem for elem in users_list if elem[1] == login]


def hashUsersPassword(password, salt):
    key = hashlib.pbkdf2_hmac(
        'sha256',
        password.encode('utf-8'),
        salt,
        100000,
        dklen=128
    )
    return salt + key


def check_password(db_password, input_password):
    salt = base64.b64decode(db_password.encode("utf-8"))[:32]
    print(db_password == str(base64.b64encode(hashUsersPassword(input_password, salt)), "utf-8"), sep="\n")
    return db_password == str(base64.b64encode(hashUsersPassword(input_password, salt)), "utf-8")


class CloseDialog(QDialog, Ui_Dialog):
    def __init__(self, *args, **kwargs):
        super(CloseDialog, self).__init__(*args, **kwargs)
        self.setupUi(self)

        self.pushButton.clicked.connect(lambda: sys.exit(app.exec()))
        self.pushButton_2.clicked.connect(lambda: self.close())


class ApplicationWindow(QStackedWidget):
    def __init__(self, widgets):
        super().__init__()
        self.setFixedWidth(800)
        self.setFixedHeight(600)
        self.setWindowTitle("Тестирующая система")
        for widget in widgets:
            self.addWidget(widget)


class GenerateTestWindow(QMainWindow, Ui_GenerateTestWindow):
    def __init__(self, *args, **kwargs):
        super(GenerateTestWindow, self).__init__(*args, **kwargs)
        self.setupUi(self)
        self.comboBox.addItems(["Арифм. операции в разл. с.с.", "Переводы между с.с."])
        self.generate_button.clicked.connect(self.generate_test)
        self.returntomenu.clicked.connect(self.openmenu)

    def generate_test(self):
        generate_systems_pdf(8, ["двоичная", "четверичная", "восьмеричная", "шестнадцатеричная"])
        generate_ariphmetic_pdf("Арифметические операции", 6, ['+', '-', "*"], [2, 4, 8, 16])

    def openmenu(self):
        app_window.setCurrentWidget(menuWindow)


class StartWindow(QMainWindow, Ui_StartWindow):
    def __init__(self, *args, **kwargs):
        super(StartWindow, self).__init__(*args, **kwargs)
        self.setupUi(self)

        self.loginButton.clicked.connect(self.openLoginWindow)
        self.loginButton_2.clicked.connect(self.changeRoleWindow)

    def openLoginWindow(self):
        app_window.setCurrentWidget(loginWindow)

    def changeRoleWindow(self):
        app_window.setCurrentWidget(roleWindow)


class RegisterWindow(QMainWindow, Ui_RegisterWindow):
    def __init__(self, *args, **kwargs):
        super(RegisterWindow, self).__init__(*args, **kwargs)
        self.setupUi(self)
        self.recovery_button.linkActivated.connect(self.returntomenu)
        self.reg_button.clicked.connect(self.register_user)
        self.login_input.editingFinished.connect(self.login_message)
        self.password_repeat.editingFinished.connect(self.password_message)
        self.login_ok = False
        self.password_ok = False
        self.reg_button.setEnabled(False)

    def register_user(self):
        salt = os.urandom(32)
        password = self.password_input.text()
        add_new_user(self.login_input.text(), hashUsersPassword(password, salt))
        app_window.setCurrentWidget(loginWindow)

    def login_message(self):
        if check_user(self.login_input.text()):
            self.error_label.setText("Такой пользователь уже существует")
            self.login_ok = False
            self.reg_button.setEnabled(False)
        else:
            self.login_ok = True
            self.error_label.setText("")
            if self.password_ok:
                self.reg_button.setEnabled(True)

    def password_message(self):
        if self.password_input.text() != self.password_repeat.text():
            self.error_label.setText("Пароли не совпадают")
            self.password_ok = False
            self.reg_button.setEnabled(False)
        else:
            self.error_label.setText("")
            self.password_ok = True
            if self.login_ok:
                self.reg_button.setEnabled(True)

    def returntomenu(self):
        app_window.setCurrentWidget(startWindow)


class LoginWindow(QMainWindow, Ui_LoginWindow):
    def __init__(self, *args, **kwargs):
        super(LoginWindow, self).__init__(*args, **kwargs)
        self.setupUi(self)

        self.login_input.editingFinished.connect(self.returnedit)
        self.login_button.clicked.connect(self.openmenu)

    def returnedit(self):
        if not check_user(self.login_input.text()):
            self.error_label.setText("Такого пользователя не существует")
            self.login_button.setEnabled(False)
        else:
            self.error_label.setText("")
            self.login_button.setEnabled(True)

    def openmenu(self):
        login, password = get_user(self.login_input.text())[0][1:]
        print(login, password)
        if check_password(password, self.password_input.text()):
            app_window.setCurrentWidget(menuWindow)
        else:
            self.error_label.setText("Введён неверный пароль")


class RoleWindow(QMainWindow, Ui_RoleWindow):
    def __init__(self, *args, **kwargs):
        super(RoleWindow, self).__init__(*args, **kwargs)
        self.setupUi(self)

        self.student_button.clicked.connect(self.openRegisterWin)
        self.teacher_button.clicked.connect(self.openRegisterWin)

    def openRegisterWin(self):
        app_window.setCurrentWidget(registerWindow)


class MenuWindow(QMainWindow, Ui_MainMenu):
    def __init__(self, *args, **kwargs):
        super(MenuWindow, self).__init__(*args, **kwargs)
        self.setupUi(self)

        self.new_test.clicked.connect(self.opengeneratewindow)
        self.exit_btn.clicked.connect(self.close_app)

    def close_app(self):
        dialog = CloseDialog()
        dialog.exec()

    def opengeneratewindow(self):
        app_window.setCurrentWidget(generateTestWindow)


if __name__ == '__main__':
    create_db()
    app = QApplication(sys.argv)
    startWindow = StartWindow()
    loginWindow = LoginWindow()
    registerWindow = RegisterWindow()
    roleWindow = RoleWindow()
    menuWindow = MenuWindow()
    generateTestWindow = GenerateTestWindow()

    app_window = ApplicationWindow([startWindow,
                                    loginWindow,
                                    registerWindow,
                                    roleWindow,
                                    menuWindow,
                                    generateTestWindow])

    app_window.show()
    sys.exit(app.exec())
